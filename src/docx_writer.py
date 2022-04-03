from docx import Document
import os

class WriteDoc:
    def __init__(self, title, paragraphs, link):
        self.title = title
        self.paragraphs = paragraphs
        self.link = link
        self.document = Document()
    def write_file(self):
        self.document.add_heading(self.title, 0)
        for paragraph in self.paragraphs:
            self.document.add_paragraph(paragraph)
        self.document.add_paragraph(f"Page link: {self.link}", style="List Bullet")
        self.document.save(f"{os.path.expanduser('~/Documents')}/{self.title}.docx")
        print("Done. Check your Documents folder for the docx file")